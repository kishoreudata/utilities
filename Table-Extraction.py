# -*- coding: utf-8 -*-
"""
Created on Mon Jul 26 09:44:52 2021

@author: Gagan Nigam
"""

import cv2
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import csv
import string

import docx
from docx import Document
import os 

try:
    from PIL import Image
except ImportError:
    import Image
import pytesseract

def valid_xml_char_ordinal(c):
    codepoint = ord(c)
    # conditions ordered by presumed frequency
    return (
        0x20 <= codepoint <= 0xD7FF or
        codepoint in (0x9, 0xA, 0xD) or
        0xE000 <= codepoint <= 0xFFFD or
        0x10000 <= codepoint <= 0x10FFFF
        )

#read your file
file=r'C:\Users\Gagan Nigam\Documents\d_drive\GKN\Table-Etraction\test_documents\ocr.jpg'
#r'/Users/marius/Desktop/Masterarbeit/Medium/Medium.png'
img = cv2.imread(file,0)
img.shape

#thresholding the image to a binary image
thresh,img_bin = cv2.threshold(img,128,255,cv2.THRESH_BINARY | cv2.THRESH_OTSU)

#inverting the image 
img_bin = 255-img_bin
#cv2.imwrite('/Users/marius/Desktop/cv_inverted.png',img_bin)
cv2.imwrite(r'C:\Users\Gagan Nigam\Documents\d_drive\GKN\Table-Etraction\test_documents\converted_documents\cv_inverted.png',img_bin)
#Plotting the image to see the output
plotting = plt.imshow(img_bin,cmap='gray')
plt.show()

# countcol(width) of kernel as 100th of total width
kernel_len = np.array(img).shape[1]//100
# Defining a vertical kernel to detect all vertical lines of image 
ver_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (1, kernel_len))
# Defining a horizontal kernel to detect all horizontal lines of image
hor_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (kernel_len, 1))
# A kernel of 2x2
kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (2, 2))

#Use vertical kernel to detect and save the vertical lines in a jpg
image_1 = cv2.erode(img_bin, ver_kernel, iterations=3)
vertical_lines = cv2.dilate(image_1, ver_kernel, iterations=3)
#cv2.imwrite("/Users/marius/Desktop/vertical.jpg",vertical_lines)
cv2.imwrite(r'C:\Users\Gagan Nigam\Documents\d_drive\GKN\Table-Etraction\test_documents\converted_documents\vertical.jpg',vertical_lines)
#Plot the generated image
plotting = plt.imshow(image_1,cmap='gray')
plt.show()

#Use horizontal kernel to detect and save the horizontal lines in a jpg
image_2 = cv2.erode(img_bin, hor_kernel, iterations=3)
horizontal_lines = cv2.dilate(image_2, hor_kernel, iterations=3)
#cv2.imwrite("/Users/marius/Desktop/horizontal.jpg",horizontal_lines)
cv2.imwrite(r'C:\Users\Gagan Nigam\Documents\d_drive\GKN\Table-Etraction\test_documents\converted_documents\horizontal.jpg',horizontal_lines)
#Plot the generated image
plotting = plt.imshow(image_2,cmap='gray')
plt.show()

# Combine horizontal and vertical lines in a new third image, with both having same weight.
img_vh = cv2.addWeighted(vertical_lines, 0.5, horizontal_lines, 0.5, 0.0)
#Eroding and thesholding the image
img_vh = cv2.erode(~img_vh, kernel, iterations=2)
thresh, img_vh = cv2.threshold(img_vh,128,255, cv2.THRESH_BINARY | cv2.THRESH_OTSU)
#cv2.imwrite("/Users/marius/Desktop/img_vh.jpg", img_vh)
cv2.imwrite(r'C:\Users\Gagan Nigam\Documents\d_drive\GKN\Table-Etraction\test_documents\converted_documents\img_vh.jpg', img_vh)
bitxor = cv2.bitwise_xor(img,img_vh)
bitnot = cv2.bitwise_not(bitxor)
#Plotting the generated image
plotting = plt.imshow(bitnot,cmap='gray')
plt.show()

# Detect contours for following box detection
contours, hierarchy = cv2.findContours(img_vh, cv2.RETR_TREE, cv2.CHAIN_APPROX_SIMPLE)

def sort_contours(cnts, method="left-to-right"):
    # initialize the reverse flag and sort index
    reverse = False
    i = 0
    # handle if we need to sort in reverse
    if method == "right-to-left" or method == "bottom-to-top":
        reverse = True
    # handle if we are sorting against the y-coordinate rather than
    # the x-coordinate of the bounding box
    if method == "top-to-bottom" or method == "bottom-to-top":
        i = 1
    # construct the list of bounding boxes and sort them from top to
    # bottom
    boundingBoxes = [cv2.boundingRect(c) for c in cnts]
    (cnts, boundingBoxes) = zip(*sorted(zip(cnts, boundingBoxes),
    key=lambda b:b[1][i], reverse=reverse))
    # return the list of sorted contours and bounding boxes
    return (cnts, boundingBoxes)

# Sort all the contours by top to bottom.
contours, boundingBoxes = sort_contours(contours, method="top-to-bottom")

#Creating a list of heights for all detected boxes
heights = [boundingBoxes[i][3] for i in range(len(boundingBoxes))]

#Get mean of heights
mean = np.mean(heights)

#Create list box to store all boxes in  
box = []
# Get position (x,y), width and height for every contour and show the contour on image
for c in contours:
    x, y, w, h = cv2.boundingRect(c)
    if (w<1000 and h<500):
        image = cv2.rectangle(img,(x,y),(x+w,y+h),(0,255,0),2)
        box.append([x,y,w,h])
        
plotting = plt.imshow(image,cmap='gray')
plt.show()

#Creating two lists to define row and column in which cell is located
row=[]
column=[]
j=0

#Sorting the boxes to their respective row and column
for i in range(len(box)):    
        
    if(i==0):
        column.append(box[i])
        previous=box[i]    
    
    else:
        if(box[i][1]<=previous[1]+mean/2):
            column.append(box[i])
            previous=box[i]            
            
            if(i==len(box)-1):
                row.append(column)        
            
        else:
            row.append(column)
            column=[]
            previous = box[i]
            column.append(box[i])
            
print(column)
print(row)

#calculating maximum number of cells
countcol = 0
for i in range(len(row)):
    countcol = len(row[i])
    if countcol > countcol:
        countcol = countcol

#Retrieving the center of each column
center = [int(row[i][j][0]+row[i][j][2]/2) for j in range(len(row[i])) if row[0]]

center=np.array(center)
center.sort()
print(center)
#Regarding the distance to the columns center, the boxes are arranged in respective order

finalboxes = []
for i in range(len(row)):
    lis=[]
    for k in range(countcol):
        lis.append([])
    for j in range(len(row[i])):
        diff = abs(center-(row[i][j][0]+row[i][j][2]/4))
        minimum = min(diff)
        indexing = list(diff).index(minimum)
        lis[indexing].append(row[i][j])
    finalboxes.append(lis)

doc = docx.Document()
#from every single image-based cell/box the strings are extracted via pytesseract and stored in a list
outer=[]
for i in range(len(finalboxes)):
    for j in range(len(finalboxes[i])):
        inner=''
        if(len(finalboxes[i][j])==0):
            outer.append(' ')
        else:
            for k in range(len(finalboxes[i][j])):
                y,x,w,h = finalboxes[i][j][k][0],finalboxes[i][j][k][1], finalboxes[i][j][k][2],finalboxes[i][j][k][3]
                finalimg = bitnot[x:x+h, y:y+w]
                kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (2, 1))
                border = cv2.copyMakeBorder(finalimg,2,2,2,2, cv2.BORDER_CONSTANT,value=[255,255])
                resizing = cv2.resize(border, None, fx=2, fy=2, interpolation=cv2.INTER_CUBIC)
                dilation = cv2.dilate(resizing, kernel,iterations=1)
                erosion = cv2.erode(dilation, kernel,iterations=2)
                
                out = pytesseract.image_to_string(erosion)
                if(len(out)==0):
                    out = pytesseract.image_to_string(erosion, config='--psm 3')
                inner = inner +" "+ out
            cleaned_string = ''.join(c for c in inner if valid_xml_char_ordinal(c))
            #doc.add_paragraph(cleaned_string)
            #outer.append(inner)
            outer.append(cleaned_string)
            
      
    # checking whether the char is printable value
        
          
        # Printing the printable values 
        
            

#Creating a dataframe of the generated OCR list
arr = np.array(outer)
dataframe = pd.DataFrame(arr.reshape(len(row), countcol))
print(dataframe)

# data = dataframe
# #data = dataframe.style.set_properties(align="left")

#Converting it in a excel-file
#data.to_excel("/Users/marius/Desktop/output.xlsx")
#data.to_excel(r'C:\Users\Gagan Nigam\Documents\d_drive\GKN\Table-Etraction\test_documents\converted_documents\output.xlsx')





# open an existing document
#rootdir = r'C:\Users\Gagan Nigam\Documents\d_drive\GKN\Table-Etraction\test_documents\converted_documents'
#doc = docx.Document(os.path.join(rootdir,'test.docx'))
#doc = docx.Document( r'C:\Users\Gagan Nigam\Documents\d_drive\GKN\Table-Etraction\test_documents\converted_documents\test.docx')



# add a table to the end and create a reference variable
#extra row is so we can add the header row
#t = doc.add_table(data.shape[0]+1, data.shape[1])
t = doc.add_table(dataframe.shape[0]+1, dataframe.shape[1])
t.style = 'TableGrid'

# add the header rows.
for j in range(dataframe.shape[-1]):
    t.cell(0,j).text = str(dataframe.columns[j])

# add the rest of the data frame
for i in range(dataframe.shape[0]):
    for j in range(dataframe.shape[-1]):
        t.cell(i+1,j).text = str(dataframe.values[i,j])

# save the doc
doc.save(r'C:\Users\Gagan Nigam\Documents\d_drive\GKN\Table-Etraction\test_documents\converted_documents\test.docx')
#doc.save(os.path.join(rootdir,'test.docx'))
